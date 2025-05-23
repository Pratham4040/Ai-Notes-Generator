from flask import Blueprint ,render_template, request , flash,redirect,url_for,send_file
import google.generativeai as genai # type: ignore
from  wtforms import FileField
from flask_wtf import FlaskForm
import os, shutil


from WEBAPP import create_app
# apikey = "AIzaSyAx-i59PwQvcID6lpx7JezKHbdjjdmhnlc"
apikey = "AIzaSyAs1DBg7iJp4yd4861maQED4dhP3jSnBr4"
genai.configure(api_key=apikey)

generation_config = {
  "temperature": 1,
  "top_p": 0.95,
  "top_k": 64,
  "max_output_tokens": 8192,
  "response_mime_type": "text/plain",
}
safety_settings = [
  {
    "category": "HARM_CATEGORY_HARASSMENT",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE",
  },
  {
    "category": "HARM_CATEGORY_HATE_SPEECH",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE",
  },
  {
    "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE",
  },
  {
    "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE",
  },
]
model = genai.GenerativeModel(
  model_name="gemini-1.5-flash-latest",
  safety_settings=safety_settings,
  generation_config=generation_config,
)

views = Blueprint('views',__name__)
@views.route('/', methods=['POST','GET'])
def home():
        if request.method == 'POST':
            if os.path.isfile("C:\ai report maker\WEBAPP\a.docx"):
              os.remove("C:\ai report maker\WEBAPP\a.docx")
            global aim 
            aim = request.form.get('aim')
            global description 
            description = request.form.get('description')
            global Experiment_Title
            Experiment_Title = request.form.get('aim')
            print(aim)
            print(description)

            first = str("Topic Tittle is [")
            after_aim = str("] The Subtopics are [")
            after_description = str("] the materials that were required were [")
            # Prompt for Brief Defination
            after_material_foraim = str("] Make Me point wise notes to help me understand my Topic, containing sections as ,Brief Defination:(just give me a small defination of the topic),Subtopics(List down the Subtopics given by me and some additional if needed),One By One Subtopics Explanation (analyse the Topic Tittle and Subtopics to give the suitable Explanations in point wise manner like type with bold the subtopic name and under that name explain it point wise, do this for all the subtopics.),Right now only give me the Brief Defination without the heading")
            # Prompt for The Subtopics 
            materialsrequired = str("] Make Me point wise notes to help me understand my Topic, containing sections as ,Brief Defination:(just give me a small defination of the topic),Subtopics(List down the Subtopics given by me and some additional if needed),One By One Subtopics Explanation (analyse the Topic Tittle and Subtopics to give the suitable Explanations in point wise manner like type with bold the subtopic name and under that name explain it point wise, do this for all the subtopics.),Right now only give me the One By One Subtopics Explanation with their heading in bold remember we need to write the ecplanation under that specefic subtopic point wise just like we do in everyday notemaking")
            
            The_aimprompt = (first + aim + after_aim + description +after_description+after_material_foraim)
            The_materialprompt = first + aim + after_aim + description +after_description+materialsrequired
  
            #starting the AI
            chat_session = model.start_chat()
            response = chat_session.send_message(The_aimprompt)
            global AimbyAI 
            AimbyAI = response.text
            #print(AimbyAI)
            response = chat_session.send_message(The_materialprompt)
            global MaterialsbyAI 
            MaterialsbyAI = response.text
            #print(MaterialsbyAI)
            
            # MAKING A DOCUMENT
            import re
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH 
            from docx.oxml.ns import qn 
            # Function to add headings and formatted text
            def add_formatted_text(doc, text):
                  # Split text into lines for processing
                  lines = text.split('\n')
                  
                  for line in lines:
                      # Check for markdown heading (e.g. # Heading)
                      if line.startswith('#'):
                          heading_level = line.count('#')
                          heading_text = line.lstrip('#').strip()
                          if heading_level == 1:
                              doc.add_heading(heading_text, level=1)
                          elif heading_level == 2:
                              doc.add_heading(heading_text, level=2)
                          # You can add more heading levels as needed.
                          continue
                      
                      # Check for bold text (e.g. **bold**)
                      bold_pattern = r'\*\*(.*?)\*\*'
                      bold_matches = re.findall(bold_pattern, line)

                      # Create a paragraph for the line
                      paragraph = doc.add_paragraph()
                      normal_text = re.split(bold_pattern, line)
                      
                      for i, part in enumerate(normal_text):
                          if i % 2 == 0:  # Normal text (not bold)
                              run = paragraph.add_run(part)
                          else:  # Bold text
                              run = paragraph.add_run(part)
                              run.bold = True

            Document = Document()
            Style = Document.styles['Normal']
            Style.font.name = 'Times New Roman'
            Style.font.size = Pt(14)
            Heading = Document.add_heading(Experiment_Title, 0)
            tittle_style = Heading.style
            tittle_style.font.name = 'Times New Roman'
            tittle_style.font.size = Pt(26)
            Heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rFonts = tittle_style.element.rPr.rFonts
            rFonts.set(qn("w:asciiTheme"), "Times New Roman")
            # FOR AIM
            HAim=Document.add_heading("Brief Description",1)
            tittle_style = HAim.style
            tittle_style.font.name = 'Times New Roman'
            tittle_style.font.size = Pt(16)
            HAim.alignment = WD_ALIGN_PARAGRAPH.CENTER
            Para = Document.add_paragraph(AimbyAI)
            Para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            #for MATERIALS REQUIRED
            HMatreq=Document.add_heading("Subtopics",1)
            tittle_style = HMatreq.style
            tittle_style.font.name = 'Times New Roman'
            tittle_style.font.size = Pt(16)
            HMatreq.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_text(Document, MaterialsbyAI)
            filepath = "C:\\ai report maker\\Ai-Report-Maker-\\WEBAPP\\notes.docx"
            Document.save(filepath)            
            return render_template("result.html")
        return render_template("home.html")


@views.route('/result', methods=['POST','GET'])
def result():
  return render_template("result.html")


@views.route('/download')
def download():
        path =  "notes.docx"
        return send_file(path, as_attachment=True)       